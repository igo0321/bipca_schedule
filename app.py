import streamlit as st
import pandas as pd
import io
import zipfile
import json
import re
import os
import copy
import smtplib
from collections import Counter
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from email.header import Header
from datetime import datetime, timedelta
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt

# ---------------------------------------------------------
# 1. ãƒ¦ãƒ¼ãƒ†ã‚£ãƒªãƒ†ã‚£
# ---------------------------------------------------------

def parse_jp_time_to_seconds(time_str):
    if not time_str:
        return 0
    s = str(time_str)
    minutes = re.search(r'(\d+)\s*[åˆ†m]', s)
    seconds = re.search(r'(\d+)\s*[ç§’s]', s)
    
    total_sec = 0
    if minutes:
        total_sec += int(minutes.group(1)) * 60
    if seconds:
        total_sec += int(seconds.group(1))
    return total_sec

def format_seconds_to_jp_label(total_seconds):
    if total_seconds <= 0:
        return "0åˆ†"
    
    minutes = total_seconds // 60
    remainder_seconds = total_seconds % 60
    
    if remainder_seconds >= 30:
        minutes += 1
        
    h = minutes // 60
    m = minutes % 60
    
    if h > 0:
        return f"{h}æ™‚é–“{m}åˆ†"
    else:
        return f"{m}åˆ†"

def format_time_label(text):
    if not text:
        return ""
    matches = re.findall(r'(\d{1,2})[:ï¼š](\d{2})', str(text))
    if len(matches) >= 2:
        start_time = f"{matches[0][0]}æ™‚{matches[0][1]}åˆ†"
        end_time = f"{matches[1][0]}æ™‚{matches[1][1]}åˆ†"
        return f"{start_time}ï½{end_time}"
    else:
        return text

def format_single_time_label(text):
    if not text:
        return ""
    match = re.search(r'(\d{1,2})[:ï¼š](\d{2})', str(text))
    if match:
        return f"{match.group(1)}æ™‚{match.group(2)}åˆ†"
    return text

def calculate_next_day_morning(date_str):
    if not date_str:
        return ""
    match = re.search(r'(\d{4})[^\d](\d{1,2})[^\d](\d{1,2})', str(date_str))
    if match:
        try:
            year, month, day = map(int, match.groups())
            dt = datetime(year, month, day)
            next_day = dt + timedelta(days=1)
            return next_day.strftime(f"%Yå¹´%mæœˆ%dæ—¥10æ™‚00åˆ†")
        except:
            return ""
    return ""

def resolve_participants_from_string(input_str, all_data_list):
    if not input_str:
        return []

    id_map = {str(item['no']): i for i, item in enumerate(all_data_list)}
    resolved_members = []
    
    parts = [p.strip() for p in input_str.replace('ã€', ',').split(',')]
    
    for part in parts:
        if not part:
            continue
        if '-' in part:
            range_parts = part.split('-')
            if len(range_parts) == 2:
                start_id = range_parts[0].strip()
                end_id = range_parts[1].strip()
                if start_id in id_map and end_id in id_map:
                    s_idx = id_map[start_id]
                    e_idx = id_map[end_id]
                    if s_idx > e_idx:
                        s_idx, e_idx = e_idx, s_idx
                    for i in range(s_idx, e_idx + 1):
                        resolved_members.append(all_data_list[i])
        else:
            if part in id_map:
                idx = id_map[part]
                resolved_members.append(all_data_list[idx])
    return resolved_members

# --- Wordæ“ä½œç³» ---

def replace_text_smart(paragraph, replacements):
    full_text = paragraph.text
    if not any(key in full_text for key in replacements):
        return

    if paragraph.runs:
        for run in paragraph.runs:
            for key, val in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, str(val))

    full_text_new = paragraph.text
    remaining_keys = [k for k in replacements if k in full_text_new]

    if remaining_keys:
        current_text = full_text_new
        for k in remaining_keys:
            current_text = current_text.replace(k, str(replacements[k]))
        
        for run in paragraph.runs:
            run.text = ""
        
        if paragraph.runs:
            paragraph.runs[0].text = current_text
        else:
            paragraph.add_run(current_text)

def fill_row_data(row, data_dict):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            replace_text_smart(paragraph, data_dict)

def replace_text_in_document_full(doc, replacements):
    for paragraph in doc.paragraphs:
        replace_text_smart(paragraph, replacements)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_smart(paragraph, replacements)
                    
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    replace_text_smart(paragraph, replacements)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_smart(paragraph, replacements)
        
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    replace_text_smart(paragraph, replacements)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_smart(paragraph, replacements)

# ---------------------------------------------------------
# 2. ãƒ¡ãƒ¼ãƒ«é€ä¿¡æ©Ÿèƒ½
# ---------------------------------------------------------

def send_email_callback():
    if 'zip_buffer' not in st.session_state or not st.session_state['zip_buffer']:
        return

    try:
        smtp_server = st.secrets["email"]["smtp_server"]
        smtp_port = st.secrets["email"]["smtp_port"]
        sender_email = st.secrets["email"]["sender_email"]
        password = st.secrets["email"]["sender_password"]
    except Exception:
        try:
            smtp_server = st.secrets["smtp"]["server"]
            smtp_port = st.secrets["smtp"]["port"]
            sender_email = st.secrets["smtp"]["sender_email"]
            password = st.secrets["smtp"]["password"]
        except:
            return

    contest_name = st.session_state.get('contest_name', 'ç„¡é¡Œ')
    user_email = st.session_state.get('user_email', 'ä¸æ˜ãªãƒ¦ãƒ¼ã‚¶ãƒ¼')
    
    file_list_str = ""
    try:
        current_pos = st.session_state['zip_buffer'].tell()
        st.session_state['zip_buffer'].seek(0)
        
        with zipfile.ZipFile(st.session_state['zip_buffer'], 'r') as zf_read:
            for name in zf_read.namelist():
                file_list_str += f"ãƒ»{name}\n"
        
        st.session_state['zip_buffer'].seek(current_pos)
    except Exception as e:
        file_list_str = f"ï¼ˆãƒ•ã‚¡ã‚¤ãƒ«ä¸€è¦§å–å¾—ã‚¨ãƒ©ãƒ¼: {e}ï¼‰"

    jst_now = datetime.utcnow() + timedelta(hours=9)
    timestamp = jst_now.strftime("%Yå¹´%mæœˆ%dæ—¥%Hæ™‚%Måˆ†")

    subject = f"æ¡ç‚¹è¡¨ç­‰ã‚’ä½œæˆã—ã¾ã—ãŸï¼š{contest_name}"
    body = f"""{user_email}ãŒä»¥ä¸‹ã®ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ç”Ÿæˆã—ã¾ã—ãŸã€‚

{file_list_str}
ç”Ÿæˆæ—¥æ™‚ï¼š{timestamp}"""
    
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = sender_email
    msg['Subject'] = Header(subject, 'utf-8')
    msg.attach(MIMEText(body, 'plain'))

    part = MIMEBase('application', 'octet-stream')
    part.set_payload(st.session_state['zip_buffer'].getvalue())
    encoders.encode_base64(part)
    
    filename = f"{contest_name}.zip"
    encoded_filename = Header(filename, 'utf-8').encode()
    part.add_header('Content-Disposition', 'attachment', filename=encoded_filename)
    
    msg.attach(part)

    try:
        server = smtplib.SMTP_SSL(smtp_server, smtp_port)
        server.login(sender_email, password)
        server.send_message(msg)
        server.quit()
    except Exception as e:
        print(f"Failed to send email: {e}")

# ---------------------------------------------------------
# 3. ãƒ‰ã‚­ãƒ¥ãƒ¡ãƒ³ãƒˆç”Ÿæˆãƒ­ã‚¸ãƒƒã‚¯
# ---------------------------------------------------------

def generate_word_from_template(template_path_or_file, groups, all_data, global_context):
    doc = Document(template_path_or_file)
    
    global_replacements = {}
    for k, v in global_context.items():
        global_replacements[f"{{{{ {k} }}}}"] = v
    replace_text_in_document_full(doc, global_replacements)

    target_table = None
    time_row_template = None
    data_row_template = None
    
    for table in doc.tables:
        t_row = None
        d_row = None
        for row in table.rows:
            row_text = "".join([c.text for c in row.cells])
            if "{{ time }}" in row_text:
                t_row = row
            if "{{ s.no }}" in row_text:
                d_row = row
        
        if t_row and d_row:
            target_table = table
            time_row_template = t_row
            data_row_template = d_row
            break
    
    if target_table:
        tbl = target_table._tbl
        time_tr = time_row_template._tr
        data_tr = data_row_template._tr
        
        tbl.remove(time_tr)
        tbl.remove(data_tr)
        
        for group in groups:
            new_tr_time = copy.deepcopy(time_tr)
            tbl.append(new_tr_time)
            new_time_row = target_table.rows[-1]
            
            raw_time = group['time_str']
            formatted_time = format_time_label(raw_time)
            fill_row_data(new_time_row, {'{{ time }}': formatted_time})

            target_members = resolve_participants_from_string(group['member_input'], all_data)
            
            for member in target_members:
                new_tr_data = copy.deepcopy(data_tr)
                tbl.append(new_tr_data)
                new_data_row = target_table.rows[-1]
                
                replacements = {
                    '{{ s.no }}': member['no'],
                    '{{ s.name }}': member['name'],
                    '{{ s.kana }}': member.get('kana', ''),
                    '{{ s.age }}': member.get('age', ''),
                    '{{ s.tel }}': member.get('tel', ''),
                    '{{ s.song }}': member['song'],
                }
                fill_row_data(new_data_row, replacements)

    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer

def generate_web_program_doc(template_path_or_file, groups, all_data, global_context):
    doc = Document(template_path_or_file)
    
    global_replacements = {}
    for k, v in global_context.items():
        global_replacements[f"{{{{ {k} }}}}"] = v
    
    replace_text_in_document_full(doc, global_replacements)
    
    bold_target_values = [
        global_context.get('contest_name', ''),
        global_context.get('contest_date', ''),
        global_context.get('contest_hall', '')
    ]
    
    def apply_bold_to_targets(doc_obj, target_values):
        def _process_para(para):
            for run in para.runs:
                for val in target_values:
                    if val and val in run.text:
                        run.font.bold = True

        for p in doc_obj.paragraphs: _process_para(p)
        for t in doc_obj.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs: _process_para(p)
    
    apply_bold_to_targets(doc, bold_target_values)

    template_time_para = None
    template_data_table = None
    
    for p in doc.paragraphs:
        if "{{ time }}" in p.text:
            template_time_para = p
            break
            
    if template_time_para:
        for table in doc.tables:
            txt = ""
            for r in table.rows:
                for c in r.cells:
                    txt += c.text
            if "{{ s.no }}" in txt:
                template_data_table = table
                break
        
        if template_data_table:
            template_p_xml = copy.deepcopy(template_time_para._p)
            template_tbl_xml = copy.deepcopy(template_data_table._tbl)
            
            parent_body = template_time_para._element.getparent()
            if parent_body is not None: parent_body.remove(template_time_para._p)
            
            parent_tbl = template_data_table._tbl.getparent()
            if parent_tbl is not None: parent_tbl.remove(template_data_table._tbl)
            
            data_tr_list = []
            header_tr_list = []
            temp_rows = list(template_tbl_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'))
            start_index = -1
            rows_per_entry = 2 

            for i, tr in enumerate(temp_rows):
                text_content = "".join([t.text for t in tr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')])
                if "{{ s.no }}" in text_content:
                    start_index = i
                    break
                else:
                    header_tr_list.append(tr)

            if start_index != -1:
                end_index = min(start_index + rows_per_entry, len(temp_rows))
                data_tr_list = temp_rows[start_index : end_index]
            
            for tr in temp_rows: template_tbl_xml.remove(tr)
            
            doc_body = doc._body._element
            
            for group in groups:
                new_p_xml = copy.deepcopy(template_p_xml)
                doc_body.append(new_p_xml)
                new_para = Paragraph(new_p_xml, doc._body)
                raw_time = group['time_str']
                formatted_time = format_time_label(raw_time)
                replace_text_smart(new_para, {'{{ time }}': formatted_time})
                
                new_tbl_xml = copy.deepcopy(template_tbl_xml)
                doc_body.append(new_tbl_xml)
                for h_tr in header_tr_list: new_tbl_xml.append(copy.deepcopy(h_tr))
                
                target_members = resolve_participants_from_string(group['member_input'], all_data)
                
                for member in target_members:
                    for tr_template in data_tr_list:
                        new_tr = copy.deepcopy(tr_template)
                        new_tbl_xml.append(new_tr)
                        
                        current_table = doc.tables[-1] 
                        current_row = current_table.rows[-1]
                        
                        for cell in current_row.cells:
                            cell_text = cell.text
                            
                            if "{{ s.no }}" in cell_text:
                                cell.text = "" 
                                p = cell.paragraphs[0]
                                run = p.add_run(f"{member['no']}")
                                run.font.bold = True
                                
                            if "{{ s.name }}" in cell_text:
                                cell.text = ""
                                p = cell.paragraphs[0]
                                run_name = p.add_run(f"{member['name']}")
                                run_name.font.bold = True
                                run_sep1 = p.add_run(" ï¼ˆ")
                                run_sep1.font.bold = False
                                if member.get('kana'):
                                    run_kana = p.add_run(f"{member['kana']}")
                                    run_kana.font.bold = False
                                run_sep2 = p.add_run("ãƒ»")
                                run_sep2.font.bold = False
                                run_age = p.add_run(f"{member.get('age', '')}")
                                run_age.font.bold = False
                                run_sep3 = p.add_run("æ­³ï¼‰")
                                run_sep3.font.bold = False
                                
                            if "{{ s.song }}" in cell_text:
                                cell.text = ""
                                p = cell.paragraphs[0]
                                run_song = p.add_run(f"{member['song']}")
                                run_song.font.bold = False

                doc_body.append(copy.deepcopy(template_p_xml))
                last_p = Paragraph(doc_body[-1], doc._body)
                last_p.text = "" 

    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer

def generate_judges_list_doc(template_path_or_file, judges_list, global_context):
    doc = Document(template_path_or_file)
    global_replacements = {}
    for k, v in global_context.items():
        global_replacements[f"{{{{ {k} }}}}"] = v
    replace_text_in_document_full(doc, global_replacements)

    for table in doc.tables:
        target_row_idx = -1
        for i, row in enumerate(table.rows):
            row_text = "".join([c.text for c in row.cells])
            if "{{ judge_name }}" in row_text:
                target_row_idx = i
                break
        
        if target_row_idx != -1:
            template_row = table.rows[target_row_idx]
            tbl = table._tbl
            tr_xml = template_row._tr
            tbl.remove(tr_xml)
            for judge in judges_list:
                new_tr = copy.deepcopy(tr_xml)
                tbl.append(new_tr)
                new_row = table.rows[-1]
                fill_row_data(new_row, {'{{ judge_name }}': judge})
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            return output_buffer

    target_para = None
    for para in doc.paragraphs:
        if "{{ judge_name }}" in para.text:
            target_para = para
            break
            
    if target_para:
        p_element = target_para._p
        parent = target_para._parent
        template_p_xml = copy.deepcopy(p_element)
        
        if hasattr(parent, '_element'):
             try: parent._element.remove(p_element)
             except: pass
        else:
             try: doc._body._body.remove(p_element)
             except: pass
        
        for judge in judges_list:
            new_p_xml = copy.deepcopy(template_p_xml)
            doc._body._body.append(new_p_xml)
            new_para = Paragraph(new_p_xml, parent)
            replace_text_smart(new_para, {'{{ judge_name }}': judge})

    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer

# ---------------------------------------------------------
# 4. è¨­å®šãƒ­ãƒ¼ãƒ‰ç”¨é–¢æ•°
# ---------------------------------------------------------

def load_settings_from_json(json_data):
    """
    JSONãƒ‡ãƒ¼ã‚¿ã‚’èª­ã¿è¾¼ã¿ã€Streamlitã®Session Stateã«åæ˜ ã•ã›ã‚‹ã€‚
    Config Versionã‚’ã‚¤ãƒ³ã‚¯ãƒªãƒ¡ãƒ³ãƒˆã™ã‚‹ã“ã¨ã§ã€ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆã®å¼·åˆ¶ãƒªãƒ•ãƒ¬ãƒƒã‚·ãƒ¥ã‚’è¡Œã†ã€‚
    """
    # 1. åŸºæœ¬ãƒ‡ãƒ¼ã‚¿
    if 'groups' in json_data:
        st.session_state['groups'] = json_data['groups']
    if 'judges' in json_data:
        st.session_state['judges'] = json_data['judges']
    if 'contest_name' in json_data:
        st.session_state['contest_name'] = json_data['contest_name']
    
    # 2. è©³ç´°è¨­å®š
    if 'contest_details' in json_data:
        st.session_state['contest_details'] = json_data['contest_details']
    
    # 3. Excelè¨­å®š (å¾Œã§Excelãƒ­ãƒ¼ãƒ‰æ™‚ã«ä½¿ç”¨ã™ã‚‹ãŸã‚ä¿å­˜)
    if 'excel_config' in json_data:
        st.session_state['saved_excel_config'] = json_data['excel_config']
    
    # 4. ãƒãƒ¼ã‚¸ãƒ§ãƒ³æ›´æ–°ï¼ˆã“ã‚Œã«ã‚ˆã‚Šã€å…¨ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆã®keyãŒå¤‰ã‚ã‚Šã€å€¤ãŒå†èª­è¾¼ã•ã‚Œã‚‹ï¼‰
    st.session_state['config_version'] += 1

# ---------------------------------------------------------
# 5. ãƒ¡ã‚¤ãƒ³ã‚¢ãƒ—ãƒªã‚±ãƒ¼ã‚·ãƒ§ãƒ³UI
# ---------------------------------------------------------
def main():
    st.set_page_config(layout="wide", page_title="ã‚³ãƒ³ã‚¯ãƒ¼ãƒ«è³‡æ–™ä½œæˆ")
    
    # åˆæœŸåŒ–
    if 'config_version' not in st.session_state:
        st.session_state['config_version'] = 0
    if 'last_loaded_json_name' not in st.session_state:
        st.session_state['last_loaded_json_name'] = None
    if 'groups' not in st.session_state:
        st.session_state['groups'] = [{'member_input': '', 'time_str': '13:00-14:10'}]
    if 'judges' not in st.session_state:
        st.session_state['judges'] = ["å¯©æŸ»å“¡A"]
    if 'saved_excel_config' not in st.session_state:
        st.session_state['saved_excel_config'] = None
    if 'contest_details' not in st.session_state:
        st.session_state['contest_details'] = {
            'date': '', 'hall': '', 'open': '10:00', 'reception': '10:45-15:30',
            'start': '11:00', 'end': '14:00', 'result': '', 'method': 'å…¬å¼ã‚µã‚¤ãƒˆä¸Šã§æ²è¼‰'
        }

    # --- 0. ãƒ¡ãƒ¼ãƒ«ã‚¢ãƒ‰ãƒ¬ã‚¹ç¢ºèª (Gateway) ---
    if 'user_email' not in st.session_state:
        st.session_state['user_email'] = None

    if not st.session_state['user_email']:
        st.title("ğŸ¹ ã‚³ãƒ³ã‚¯ãƒ¼ãƒ«é‹å–¶è³‡æ–™ã‚¸ã‚§ãƒãƒ¬ãƒ¼ã‚¿ãƒ¼")
        st.info("ä½¿ç”¨å±¥æ­´ã‚’ç¢ºèªã™ã‚‹ãŸã‚ã€ãƒ¡ãƒ¼ãƒ«ã‚¢ãƒ‰ãƒ¬ã‚¹ã®å…¥åŠ›ã‚’ãŠé¡˜ã„ã—ã¾ã™ã€‚")
        
        with st.form("email_login_form"):
            input_email = st.text_input("ã”æ‹…å½“è€…æ§˜ ãƒ¡ãƒ¼ãƒ«ã‚¢ãƒ‰ãƒ¬ã‚¹", placeholder="example@example.com")
            submit_login = st.form_submit_button("åˆ©ç”¨ã‚’é–‹å§‹ã™ã‚‹")
            
            if submit_login:
                if input_email and "@" in input_email:
                    st.session_state['user_email'] = input_email
                    st.rerun()
                else:
                    st.error("æœ‰åŠ¹ãªãƒ¡ãƒ¼ãƒ«ã‚¢ãƒ‰ãƒ¬ã‚¹ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚")
        st.stop()

    # --- ä»¥ä¸‹ã€ãƒ¡ã‚¤ãƒ³ã‚³ãƒ³ãƒ†ãƒ³ãƒ„ ---
    st.title("ğŸ¹ ã‚³ãƒ³ã‚¯ãƒ¼ãƒ«é‹å–¶è³‡æ–™ã‚¸ã‚§ãƒãƒ¬ãƒ¼ã‚¿ãƒ¼ (Wordç‰ˆ)")
    st.markdown(f"**ãƒ­ã‚°ã‚¤ãƒ³ä¸­:** {st.session_state['user_email']}")
    
    # ãƒãƒ¼ã‚¸ãƒ§ãƒ³ç•ªå·ã®å–å¾—ï¼ˆã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆKeyç”Ÿæˆç”¨ï¼‰
    ver = st.session_state['config_version']

    # --- ã‚µã‚¤ãƒ‰ãƒãƒ¼: è¨­å®šèª­ã¿è¾¼ã¿ ---
    with st.sidebar:
        st.header("âš™ï¸ è¨­å®šç®¡ç†")
        
        # ã‚­ãƒ¼ã¯å›ºå®šã—ã€èª­ã¿è¾¼ã¿å‡¦ç†ã§rerunã‚’ä½¿ã‚ãªã„æ–¹å¼ã«å¤‰æ›´
        uploaded_config = st.file_uploader(
            "è¨­å®šãƒ•ã‚¡ã‚¤ãƒ«(JSON)ã‚’èª­ã¿è¾¼ã‚€", 
            type=['json'], 
            key="json_config_uploader_fixed" 
        )

        if uploaded_config:
            if uploaded_config.name != st.session_state['last_loaded_json_name']:
                try:
                    # æ˜ç¤ºçš„ã«UTF-8ã§èª­ã¿è¾¼ã‚€
                    content = uploaded_config.getvalue().decode("utf-8")
                    config_data = json.loads(content)
                    
                    # è¨­å®šãƒ­ãƒ¼ãƒ‰ & ãƒãƒ¼ã‚¸ãƒ§ãƒ³ã‚¢ãƒƒãƒ—
                    load_settings_from_json(config_data)
                    
                    # èª­ã¿è¾¼ã¿æ¸ˆã¿ãƒ•ãƒ©ã‚°æ›´æ–°
                    st.session_state['last_loaded_json_name'] = uploaded_config.name
                    st.success("è¨­å®šã‚’èª­ã¿è¾¼ã¿ã¾ã—ãŸã€‚")
                    
                    # ã“ã“ã§ st.rerun() ã¯å‘¼ã°ãªã„ï¼
                    # verãŒå¢—ãˆãŸã®ã§ã€ã“ã®å¾Œã®å‡¦ç†ã§è‡ªå‹•çš„ã«æ–°ã—ã„ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆãŒç”Ÿæˆã•ã‚Œã€å€¤ãŒåæ˜ ã•ã‚Œã‚‹ã€‚
                    
                except Exception as e:
                    st.error(f"è¨­å®šèª­ã¿è¾¼ã¿ã‚¨ãƒ©ãƒ¼: {e}")
        else:
            st.session_state['last_loaded_json_name'] = None

    # --- 1. åç°¿ãƒ‡ãƒ¼ã‚¿ (Excel) ---
    st.header("1. åç°¿ãƒ‡ãƒ¼ã‚¿ (Excel)")
    
    uploaded_excel = st.file_uploader(
        "åç°¿Excelãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰", 
        type=['xlsx', 'xls', 'csv'], 
        key="excel_uploader_fixed"
    )
    
    all_data = []
    excel_config_to_save = {}
    
    if uploaded_excel:
        try:
            # --- ã‚·ãƒ¼ãƒˆé¸æŠãƒ­ã‚¸ãƒƒã‚¯ (ä¿å­˜ã•ã‚ŒãŸè¨­å®šã®å„ªå…ˆä½¿ç”¨) ---
            saved_config = st.session_state.get('saved_excel_config', {})
            saved_sheet = saved_config.get('sheet_name') if saved_config else None
            
            df = None
            selected_sheet = None
            
            if uploaded_excel.name.endswith('.csv'):
                df = pd.read_csv(uploaded_excel)
                selected_sheet = "CSV"
            else:
                xls = pd.ExcelFile(uploaded_excel)
                sheet_names = xls.sheet_names
                
                # ãƒ‡ãƒ•ã‚©ãƒ«ãƒˆã‚¤ãƒ³ãƒ‡ãƒƒã‚¯ã‚¹ã®æ±ºå®š
                default_sheet_idx = 0
                if saved_sheet and saved_sheet in sheet_names:
                    default_sheet_idx = sheet_names.index(saved_sheet)
                
                # ã‚»ãƒ¬ã‚¯ãƒˆãƒœãƒƒã‚¯ã‚¹ã¯ãƒãƒ¼ã‚¸ãƒ§ãƒ³ä¾å­˜ã§å†ç”Ÿæˆã•ã›ã‚‹ (è¨­å®šåæ˜ ã®ãŸã‚)
                selected_sheet = st.selectbox("ã‚·ãƒ¼ãƒˆã‚’é¸æŠ", sheet_names, index=default_sheet_idx, key=f"sheet_sel_{ver}")
                df = pd.read_excel(uploaded_excel, sheet_name=selected_sheet)

            # Excelè¨­å®šä¿å­˜ç”¨
            excel_config_to_save['sheet_name'] = selected_sheet

            # --- åˆ—ã®å‰²ã‚Šå½“ã¦ãƒ­ã‚¸ãƒƒã‚¯ (ä¿å­˜ã•ã‚ŒãŸè¨­å®šã®å„ªå…ˆä½¿ç”¨) ---
            cols = df.columns.tolist()
            
            def get_col_index(saved_key, default_heuristic_cols, all_cols, fallback_index=0):
                # 1. ä¿å­˜ã•ã‚ŒãŸè¨­å®šãŒã‚ã‚Œã°ãã‚Œã‚’ä½¿ã†
                if saved_config and saved_key in saved_config:
                    val = saved_config[saved_key]
                    if val in all_cols:
                        return all_cols.index(val)
                # 2. ãƒ’ãƒ¥ãƒ¼ãƒªã‚¹ãƒ†ã‚£ãƒƒã‚¯
                for h in default_heuristic_cols:
                    if h in all_cols:
                        return all_cols.index(h)
                # 3. ãƒ•ã‚©ãƒ¼ãƒ«ãƒãƒƒã‚¯
                return fallback_index

            c1, c2, c3, c4 = st.columns(4)
            
            # å„ã‚»ãƒ¬ã‚¯ãƒˆãƒœãƒƒã‚¯ã‚¹ã¯ {ver} ã‚’ã‚­ãƒ¼ã«å«ã‚ã‚‹ã“ã¨ã§ã€JSONãƒ­ãƒ¼ãƒ‰æ™‚ã«åˆæœŸé¸æŠä½ç½®ã‚’å†è¨ˆç®—ã•ã›ã‚‹
            idx_no = get_col_index('col_no', ["å‡ºå ´ç•ªå·", "No", "No."], cols, 0)
            col_no = c1.selectbox("å‡ºå ´ç•ªå·", cols, index=idx_no, key=f"c_no_{ver}")

            idx_name = get_col_index('col_name', ["æ°å", "åå‰"], cols, 0)
            col_name = c2.selectbox("æ°å", cols, index=idx_name, key=f"c_name_{ver}")
            
            kana_options = ["(ãªã—)"] + cols
            idx_kana = 0
            if saved_config and 'col_kana' in saved_config:
                if saved_config['col_kana'] in kana_options:
                    idx_kana = kana_options.index(saved_config['col_kana'])
            elif "ãƒ•ãƒªã‚¬ãƒŠ" in cols:
                idx_kana = cols.index("ãƒ•ãƒªã‚¬ãƒŠ") + 1
            col_kana = c3.selectbox("ãƒ•ãƒªã‚¬ãƒŠ (ä»»æ„)", kana_options, index=idx_kana, key=f"c_kana_{ver}")
            
            idx_song = get_col_index('col_song', ["æ¼”å¥æ›²ç›®", "æ›²ç›®"], cols, 0)
            col_song = c4.selectbox("æ¼”å¥æ›²ç›®", cols, index=idx_song, key=f"c_song_{ver}")
            
            c5, c6, c7 = st.columns(3)
            
            age_options = ["(ãªã—)"] + cols
            idx_age = 0
            if saved_config and 'col_age' in saved_config:
                 if saved_config['col_age'] in age_options: idx_age = age_options.index(saved_config['col_age'])
            elif "å¹´é½¢" in cols:
                 idx_age = cols.index("å¹´é½¢") + 1
            col_age = c5.selectbox("å¹´é½¢åˆ— (ä»»æ„)", age_options, index=idx_age, key=f"c_age_{ver}")

            tel_options = ["(ãªã—)"] + cols
            idx_tel = 0
            if saved_config and 'col_tel' in saved_config:
                 if saved_config['col_tel'] in tel_options: idx_tel = tel_options.index(saved_config['col_tel'])
            elif "é›»è©±ç•ªå·" in cols:
                 idx_tel = cols.index("é›»è©±ç•ªå·") + 1
            col_tel = c6.selectbox("é›»è©±ç•ªå·åˆ— (å—ä»˜è¡¨ç”¨)", tel_options, index=idx_tel, key=f"c_tel_{ver}")

            dur_options = ["(ãªã—)"] + cols
            idx_dur = 0
            if saved_config and 'col_duration' in saved_config:
                 if saved_config['col_duration'] in dur_options: idx_dur = dur_options.index(saved_config['col_duration'])
            elif "æ¼”å¥æ™‚é–“" in cols:
                 idx_dur = cols.index("æ¼”å¥æ™‚é–“") + 1
            col_duration = c7.selectbox("æ¼”å¥æ™‚é–“åˆ— (è‡ªå‹•è¨ˆç®—ç”¨)", dur_options, index=idx_dur, key=f"c_dur_{ver}")

            excel_config_to_save.update({
                'col_no': col_no,
                'col_name': col_name,
                'col_kana': col_kana,
                'col_song': col_song,
                'col_age': col_age,
                'col_tel': col_tel,
                'col_duration': col_duration
            })

            st.markdown("---")

            for _, row in df.iterrows():
                kana_val = str(row[col_kana]) if col_kana != "(ãªã—)" else ""
                age_val = str(row[col_age]) if col_age != "(ãªã—)" else ""
                tel_val = str(row[col_tel]) if col_tel != "(ãªã—)" else ""
                dur_seconds = 0
                if col_duration != "(ãªã—)":
                    raw_dur = str(row[col_duration])
                    dur_seconds = parse_jp_time_to_seconds(raw_dur)

                all_data.append({
                    'no': str(row[col_no]), 
                    'name': str(row[col_name]),
                    'kana': kana_val,
                    'song': str(row[col_song]),
                    'age': age_val,
                    'tel': tel_val,
                    'duration_sec': dur_seconds
                })
            
            st.write(f"èª­ã¿è¾¼ã¿å®Œäº†: {len(all_data)} ä»¶ã®ãƒ‡ãƒ¼ã‚¿")

            # --- 2. ãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆé¸æŠ ---
            st.header("2. Wordãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆé¸æŠ")
            
            TEMPLATE_DIR = "templates"
            template_files = []
            if os.path.exists(TEMPLATE_DIR):
                template_files = [f for f in os.listdir(TEMPLATE_DIR) if f.endswith(".docx") and not f.startswith("~$")]
            
            score_template_path = None
            reception_template_path = None
            web_template_path = None
            judges_list_template_path = None
            use_manual_upload = False

            if template_files:
                idx_score = 0
                idx_reception = 0
                idx_web = 0
                idx_judges = 0
                for i, f in enumerate(template_files):
                    if "æ¡ç‚¹è¡¨" in f: idx_score = i
                    if "å—ä»˜è¡¨" in f: idx_reception = i
                    if "WEB" in f or "ãƒ—ãƒ­ã‚°ãƒ©ãƒ " in f: idx_web = i
                    if "å¯©æŸ»å“¡" in f and "ãƒªã‚¹ãƒˆ" not in f: idx_judges = i
                
                col_t1, col_t2 = st.columns(2)
                col_t3, col_t4 = st.columns(2)
                
                with col_t1:
                    selected_score_file = st.selectbox("æ¡ç‚¹è¡¨ãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆ", template_files, index=idx_score, key=f"tpl_sc_{ver}")
                    score_template_path = os.path.join(TEMPLATE_DIR, selected_score_file)
                with col_t2:
                    selected_reception_file = st.selectbox("å—ä»˜è¡¨ãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆ", template_files, index=idx_reception, key=f"tpl_rc_{ver}")
                    reception_template_path = os.path.join(TEMPLATE_DIR, selected_reception_file)
                with col_t3:
                    selected_web_file = st.selectbox("WEBãƒ—ãƒ­ã‚°ãƒ©ãƒ ãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆ", template_files, index=idx_web, key=f"tpl_wb_{ver}")
                    web_template_path = os.path.join(TEMPLATE_DIR, selected_web_file)
                with col_t4:
                    selected_judges_file = st.selectbox("å¯©æŸ»å“¡ãƒªã‚¹ãƒˆãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆ", template_files, index=idx_judges, key=f"tpl_jd_{ver}")
                    judges_list_template_path = os.path.join(TEMPLATE_DIR, selected_judges_file)
                
                if st.checkbox("ãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆã‚’æ‰‹å‹•ã§ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ã™ã‚‹", key=f"chk_manual_{ver}"):
                    use_manual_upload = True
            else:
                st.warning("templatesãƒ•ã‚©ãƒ«ãƒ€ãŒè¦‹ã¤ã‹ã‚‰ãªã„ã‹ã€docxãƒ•ã‚¡ã‚¤ãƒ«ãŒã‚ã‚Šã¾ã›ã‚“ã€‚æ‰‹å‹•ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ãƒ¢ãƒ¼ãƒ‰ã«åˆ‡ã‚Šæ›¿ãˆã¾ã™ã€‚")
                use_manual_upload = True

            if use_manual_upload:
                c_up1, c_up2 = st.columns(2)
                c_up3, c_up4 = st.columns(2)
                uploaded_score_template = c_up1.file_uploader("æ¡ç‚¹è¡¨ãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆ (.docx)", type=['docx'], key=f"up_sc_{ver}")
                uploaded_reception_template = c_up2.file_uploader("å—ä»˜è¡¨ãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆ (.docx)", type=['docx'], key=f"up_rc_{ver}")
                uploaded_web_template = c_up3.file_uploader("WEBãƒ—ãƒ­ã‚°ãƒ©ãƒ ãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆ (.docx)", type=['docx'], key=f"up_wb_{ver}")
                uploaded_judges_template = c_up4.file_uploader("å¯©æŸ»å“¡ãƒªã‚¹ãƒˆãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆ (.docx)", type=['docx'], key=f"up_jd_{ver}")
                
                if uploaded_score_template: score_template_path = uploaded_score_template
                if uploaded_reception_template: reception_template_path = uploaded_reception_template
                if uploaded_web_template: web_template_path = uploaded_web_template
                if uploaded_judges_template: judges_list_template_path = uploaded_judges_template

            # --- 3. ã‚°ãƒ«ãƒ¼ãƒ—ãƒ»ã‚¹ã‚±ã‚¸ãƒ¥ãƒ¼ãƒ«è¨­å®š (Dynamic Key Implemented) ---
            st.header("3. ã‚°ãƒ«ãƒ¼ãƒ—ãƒ»ã‚¹ã‚±ã‚¸ãƒ¥ãƒ¼ãƒ«è¨­å®š")
            
            def add_group():
                st.session_state['groups'].append({'member_input': '', 'time_str': ''})
            def move_group_up(idx):
                if idx > 0:
                    st.session_state['groups'][idx], st.session_state['groups'][idx-1] = st.session_state['groups'][idx-1], st.session_state['groups'][idx]
            def move_group_down(idx):
                if idx < len(st.session_state['groups']) - 1:
                    st.session_state['groups'][idx], st.session_state['groups'][idx+1] = st.session_state['groups'][idx+1], st.session_state['groups'][idx]
            def remove_group(idx):
                st.session_state['groups'].pop(idx)

            st.button("ï¼‹ ã‚°ãƒ«ãƒ¼ãƒ—è¿½åŠ ", on_click=add_group, key=f"btn_add_grp_{ver}")

            for i, grp in enumerate(st.session_state['groups']):
                c_sort, c_input, c_total, c_time, c_del = st.columns([0.8, 3, 1.2, 2, 0.5])
                
                with c_sort:
                    if st.button("â–²", key=f"up_{i}_{ver}"):
                        move_group_up(i)
                        st.rerun()
                    if st.button("â–¼", key=f"down_{i}_{ver}"):
                        move_group_down(i)
                        st.rerun()

                input_val = c_input.text_input(
                    f"ã‚°ãƒ«ãƒ¼ãƒ— {i+1} å¯¾è±¡ç•ªå·",
                    value=grp['member_input'],
                    key=f"g_in_{i}_{ver}",
                    placeholder="ä¾‹: A01-A05, C01"
                )
                st.session_state['groups'][i]['member_input'] = input_val

                current_members = resolve_participants_from_string(input_val, all_data)
                total_sec = sum(m['duration_sec'] for m in current_members)
                time_display = format_seconds_to_jp_label(total_sec)
                
                with c_total:
                    st.markdown(f"""
                    <div style="margin-bottom: 0px;">
                        <label style="font-size: 14px; color: rgb(49, 51, 63); margin-bottom: 0.5rem; display: block;">
                            åˆè¨ˆæ¼”å¥æ™‚é–“
                        </label>
                        <div style="
                            background-color: rgba(28, 131, 225, 0.1); 
                            border: 1px solid rgba(28, 131, 225, 0.1);
                            border-radius: 0.5rem;
                            padding: 0px 10px;
                            min-height: 2.5rem;
                            height: auto;
                            display: flex;
                            align-items: center;
                            color: rgb(0, 66, 128);
                            font-size: 1rem;
                            line-height: 1.5;
                        ">
                            è¨ˆ: {time_display}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)

                time_val = c_time.text_input(
                    "æ™‚é–“",
                    value=grp['time_str'],
                    key=f"g_time_{i}_{ver}",
                    placeholder="ä¾‹: 13:00-14:00"
                )
                st.session_state['groups'][i]['time_str'] = time_val

                with c_del:
                    st.markdown("<div style='margin-top: 1.8rem;'></div>", unsafe_allow_html=True)
                    if st.button("Ã—", key=f"del_{i}_{ver}"):
                        remove_group(i)
                        st.rerun()

            # --- 4. å¯©æŸ»å“¡è¨­å®š (Dynamic Key Implemented) ---
            st.header("4. å¯©æŸ»å“¡è¨­å®š")
            
            def add_judge():
                st.session_state['judges'].append("")
            
            st.button("ï¼‹ å¯©æŸ»å“¡è¿½åŠ ", on_click=add_judge, key=f"btn_add_jdg_{ver}")

            # ãƒ«ãƒ¼ãƒ—ã§ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆç”Ÿæˆã€‚Keyã«verã‚’å«ã‚ã‚‹ã“ã¨ã§ã€JSONãƒ­ãƒ¼ãƒ‰æ™‚ã«å¼·åˆ¶ãƒªãƒ•ãƒ¬ãƒƒã‚·ãƒ¥
            for i in range(len(st.session_state['judges'])):
                val = st.text_input(
                    f"å¯©æŸ»å“¡ {i+1}", 
                    value=st.session_state['judges'][i], 
                    key=f"judge_input_{i}_{ver}" 
                )
                st.session_state['judges'][i] = val

            contest_name = st.text_input("ã‚³ãƒ³ã‚¯ãƒ¼ãƒ«å (ãƒ•ã‚¡ã‚¤ãƒ«åç­‰ã«ä½¿ç”¨)", 
                                         value=st.session_state['contest_name'],
                                         key=f"input_contest_name_{ver}")
            st.session_state['contest_name'] = contest_name

            # --- 5. å¯©æŸ»ä¼šè©³ç´° (Dynamic Key Implemented) ---
            st.header("5. å¯©æŸ»ä¼šè©³ç´°")
            st.info("â€»ã“ã“ã§å…¥åŠ›ã—ãŸå†…å®¹ã¯Wordå‡ºåŠ›æ™‚ã«è‡ªå‹•çš„ã«å½¢å¼å¤‰æ›ã•ã‚Œã¦æŒ¿å…¥ã•ã‚Œã¾ã™ã€‚")
            
            det_current = st.session_state['contest_details']

            def on_date_change():
                # Dynamic Keyã‹ã‚‰ç¾åœ¨ã®å€¤ã‚’å–å¾—ã—ã¦è¨ˆç®—
                v = st.session_state['config_version']
                current_date = st.session_state.get(f"detail_date_{v}", "")
                calculated = calculate_next_day_morning(current_date)
                if calculated:
                    # çµæœç™ºè¡¨æ—¥æ™‚ã‚’æ›´æ–°
                    st.session_state['contest_details']['result'] = calculated
            
            col_d1, col_d2 = st.columns(2)
            date_val = col_d1.text_input("é–‹å‚¬æ—¥æ™‚ (ä¾‹: 2025å¹´12æœˆ21æ—¥)", value=det_current['date'], key=f"detail_date_{ver}", on_change=on_date_change)
            hall_val = col_d2.text_input("ä¼šå ´", value=det_current['hall'], key=f"detail_hall_{ver}")
            
            col_d3, col_d4, col_d5, col_d6 = st.columns(4)
            open_val = col_d3.text_input("é–‹å ´æ™‚åˆ» (ä¾‹: 10:00)", value=det_current['open'], key=f"detail_open_{ver}")
            start_val = col_d4.text_input("å¯©æŸ»é–‹å§‹ (ä¾‹: 11:00)", value=det_current['start'], key=f"detail_start_{ver}")
            end_val = col_d5.text_input("å¯©æŸ»çµ‚äº† (ä¾‹: 14:00)", value=det_current['end'], key=f"detail_end_{ver}")
            reception_val = col_d6.text_input("å—ä»˜æ™‚é–“ (ä¾‹: 10:45-15:30)", value=det_current['reception'], key=f"detail_reception_{ver}")

            col_d7, col_d8 = st.columns(2)
            result_val = col_d7.text_input("çµæœç™ºè¡¨æ—¥æ™‚ (è‡ªå‹•è¨ˆç®—)", value=det_current['result'], key=f"detail_result_{ver}")
            
            method_options = ["å…¬å¼ã‚µã‚¤ãƒˆä¸Šã§æ²è¼‰", "ä¼šå ´ãƒ­ãƒ“ãƒ¼ã‚‚ã—ãã¯ãƒ›ãƒ¯ã‚¤ã‚¨ã§æ²è¼‰", "è¡¨å½°å¼ã«ã¦ç™ºè¡¨", "ãã®ä»–"]
            curr_method = det_current.get('method', "å…¬å¼ã‚µã‚¤ãƒˆä¸Šã§æ²è¼‰")
            idx_method = method_options.index(curr_method) if curr_method in method_options else 0
            method_val = col_d8.selectbox("çµæœç™ºè¡¨æ–¹å¼", method_options, index=idx_method, key=f"detail_method_{ver}")

            # æœ€æ–°ã®å…¥åŠ›å€¤ã‚’ä¿å­˜ç”¨è¾æ›¸ã«æ ¼ç´
            det_updated = {
                'date': date_val, 'hall': hall_val, 
                'open': open_val, 'start': start_val, 'end': end_val, 
                'reception': reception_val, 'result': result_val, 'method': method_val
            }
            st.session_state['contest_details'] = det_updated

            # --- 6. ãƒ•ã‚¡ã‚¤ãƒ«å‡ºåŠ› ---
            st.header("6. ãƒ•ã‚¡ã‚¤ãƒ«å‡ºåŠ›")
            if st.button("ãƒ•ã‚¡ã‚¤ãƒ«ç”Ÿæˆã‚’å®Ÿè¡Œ", type="primary", key=f"btn_gen_{ver}"):
                # ãƒãƒªãƒ‡ãƒ¼ã‚·ãƒ§ãƒ³
                assigned_nos = []
                for grp in st.session_state['groups']:
                    members = resolve_participants_from_string(grp['member_input'], all_data)
                    for m in members:
                        assigned_nos.append(m['no'])
                
                counts = Counter(assigned_nos)
                duplicates = [no for no, count in counts.items() if count > 1]
                
                if duplicates:
                    st.error(f"â›” ã‚¨ãƒ©ãƒ¼: ä»¥ä¸‹ã®å‡ºå ´ç•ªå·ãŒè¤‡æ•°ã®ã‚°ãƒ«ãƒ¼ãƒ—ã«é‡è¤‡ã—ã¦ç™»éŒ²ã•ã‚Œã¦ã„ã¾ã™ã€‚\n{', '.join(duplicates)}")
                    return 
                
                all_nos_set = set(item['no'] for item in all_data)
                assigned_nos_set = set(assigned_nos)
                unregistered = sorted(list(all_nos_set - assigned_nos_set))
                
                if unregistered:
                    st.warning(f"âš ï¸ æ³¨æ„: ä»¥ä¸‹ã®å‡ºå ´ç•ªå·ã¯ã©ã®ã‚°ãƒ«ãƒ¼ãƒ—ã«ã‚‚ç™»éŒ²ã•ã‚Œã¦ã„ã¾ã›ã‚“ã€‚\n{', '.join(unregistered)}")

                if not score_template_path:
                    st.error("æ¡ç‚¹è¡¨ãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆãŒé¸æŠã•ã‚Œã¦ã„ã¾ã›ã‚“ã€‚")
                    return
                if not web_template_path:
                    st.warning("WEBãƒ—ãƒ­ã‚°ãƒ©ãƒ ãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆãŒé¸æŠã•ã‚Œã¦ã„ã¾ã›ã‚“ã€‚")
                if not judges_list_template_path:
                    st.warning("å¯©æŸ»å“¡ãƒªã‚¹ãƒˆãƒ†ãƒ³ãƒ—ãƒ¬ãƒ¼ãƒˆãŒé¸æŠã•ã‚Œã¦ã„ã¾ã›ã‚“ã€‚")

                valid_judges = [j for j in st.session_state['judges'] if j.strip()]
                
                details_formatted = {
                    'contest_date': det_updated['date'],
                    'contest_hall': det_updated['hall'],
                    'contest_open': format_single_time_label(det_updated['open']),
                    'contest_reception': format_time_label(det_updated['reception']),
                    'contest_start': format_single_time_label(det_updated['start']),
                    'contest_end': format_single_time_label(det_updated['end']),
                    'contest_result': det_updated['result'],
                    'contest_method': det_updated['method']
                }

                # Config JSONä½œæˆ
                config_json = json.dumps({
                    'groups': st.session_state['groups'],
                    'judges': valid_judges,
                    'contest_name': contest_name,
                    'contest_details': det_updated,
                    'excel_config': excel_config_to_save
                }, ensure_ascii=False, indent=2)

                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                    
                    base_context = {
                        'contest_name': contest_name,
                        **details_formatted
                    }

                    # ç”Ÿæˆå‡¦ç†
                    for judge in valid_judges:
                        try:
                            if hasattr(score_template_path, 'seek'): score_template_path.seek(0)
                            context = base_context.copy()
                            context['judge_name'] = judge
                            doc_io = generate_word_from_template(score_template_path, st.session_state['groups'], all_data, context)
                            zf.writestr(f"æ¡ç‚¹è¡¨_{judge}.docx", doc_io.getvalue())
                        except Exception as e:
                            st.error(f"æ¡ç‚¹è¡¨ç”Ÿæˆã‚¨ãƒ©ãƒ¼ ({judge}): {e}")

                    if reception_template_path:
                        try:
                            if hasattr(reception_template_path, 'seek'): reception_template_path.seek(0)
                            context = base_context.copy()
                            context['judge_name'] = 'å—ä»˜ç”¨'
                            doc_io = generate_word_from_template(reception_template_path, st.session_state['groups'], all_data, context)
                            zf.writestr("å—ä»˜è¡¨.docx", doc_io.getvalue())
                        except Exception as e:
                            st.error(f"å—ä»˜è¡¨ç”Ÿæˆã‚¨ãƒ©ãƒ¼: {e}")

                    if web_template_path:
                        try:
                            if hasattr(web_template_path, 'seek'): web_template_path.seek(0)
                            context = base_context.copy()
                            context['judge_name'] = ''
                            doc_io = generate_web_program_doc(web_template_path, st.session_state['groups'], all_data, context)
                            zf.writestr("WEBãƒ—ãƒ­ã‚°ãƒ©ãƒ .docx", doc_io.getvalue())
                        except Exception as e:
                            st.error(f"WEBãƒ—ãƒ­ã‚°ãƒ©ãƒ ç”Ÿæˆã‚¨ãƒ©ãƒ¼: {e}")
                            
                    if judges_list_template_path:
                         try:
                            if hasattr(judges_list_template_path, 'seek'): judges_list_template_path.seek(0)
                            context = base_context.copy()
                            doc_io = generate_judges_list_doc(judges_list_template_path, valid_judges, context)
                            zf.writestr("æœ¬æ—¥ã®å¯©æŸ»å“¡.docx", doc_io.getvalue())
                         except Exception as e:
                            st.error(f"å¯©æŸ»å“¡ãƒªã‚¹ãƒˆç”Ÿæˆã‚¨ãƒ©ãƒ¼: {e}")

                    if os.path.exists(TEMPLATE_DIR):
                        pdf_files = [f for f in os.listdir(TEMPLATE_DIR) if f.endswith(".pdf")]
                        for pdf_file in pdf_files:
                            pdf_path = os.path.join(TEMPLATE_DIR, pdf_file)
                            zf.write(pdf_path, arcname=pdf_file)

                    zf.writestr("è¨­å®šãƒ‡ãƒ¼ã‚¿.json", config_json)
                
                st.session_state['zip_buffer'] = zip_buffer
                st.success("ç”Ÿæˆå®Œäº†ï¼ä¸‹ã®ãƒœã‚¿ãƒ³ã‹ã‚‰ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰ã—ã¦ãã ã•ã„ã€‚")
            
            if 'zip_buffer' in st.session_state and st.session_state['zip_buffer']:
                st.download_button(
                    label="ZIPãƒ•ã‚¡ã‚¤ãƒ«ã‚’ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰",
                    data=st.session_state['zip_buffer'].getvalue(),
                    file_name=f"{contest_name}.zip",
                    mime="application/zip",
                    on_click=send_email_callback,
                    key=f"dl_btn_{ver}"
                )

        except Exception as e:
            st.error(f"ã‚¨ãƒ©ãƒ¼ãŒç™ºç”Ÿã—ã¾ã—ãŸ: {e}")

if __name__ == "__main__":
    main()
